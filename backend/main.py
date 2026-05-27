from pathlib import Path
import os
import json
import requests
import re
import shutil
from typing import List, Optional

from docx import Document
from dotenv import load_dotenv
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from pypdf import PdfReader


load_dotenv(dotenv_path=Path(__file__).with_name(".env"))

app = FastAPI(title="AI Teaching Material Audit Backend")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)


def build_empty_course_standard_info() -> dict:
    return {
        "course_objectives": [],
        "teaching_content": [],
        "assessment_methods": [],
        "raw_sections": {
            "objective_section": "",
            "content_section": "",
            "assessment_section": "",
        },
    }


def safe_filename(filename: Optional[str]) -> str:
    if not filename:
        return "uploaded_file"

    clean_name = Path(filename).name.strip()
    return clean_name or "uploaded_file"


def extract_text_from_docx(file_path: Path) -> str:
    doc = Document(file_path)
    texts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            texts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                texts.append(" | ".join(row_texts))

    return "\n".join(texts)


def extract_text_from_pdf(file_path: Path) -> str:
    reader = PdfReader(str(file_path))
    texts = []

    for page in reader.pages:
        page_text = page.extract_text() or ""
        page_text = page_text.strip()
        if page_text:
            texts.append(page_text)

    return "\n".join(texts)


def extract_text(file_path: Path) -> str:
    suffix = file_path.suffix.lower()

    if suffix == ".docx":
        return extract_text_from_docx(file_path)

    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)

    return ""


def normalize_course_name(name: str) -> str:
    if not name:
        return ""

    name = name.strip()
    name = name.replace("《", "").replace("》", "")
    name = name.replace("“", "").replace("”", "")
    name = name.replace('"', "").replace("'", "")
    name = re.sub(r"(课程标准|课程教学大纲|教学大纲|课程大纲)$", "", name)

    return name.strip()


def guess_course_name_from_text(text: str) -> str:
    if not text:
        return ""

    head = text[:3000]
    patterns = [
        r"课程名称\s*[:：]\s*([^\n\r|，。；;]+)",
        r"课程名称\s+([^\n\r|，。；;]+)",
        r"《([^》]{2,40})》\s*课程标准",
        r"([^\n\r]{2,40})\s*课程标准",
    ]

    for pattern in patterns:
        match = re.search(pattern, head)
        if match:
            course_name = normalize_course_name(match.group(1))
            if 2 <= len(course_name) <= 30:
                return course_name

    return ""


def guess_course_name_from_filename(filename: str) -> str:
    name = filename or ""
    name = re.sub(r"\.(docx|doc|pdf|pptx|ppt|xlsx|xls)$", "", name, flags=re.I)
    name = re.sub(
        r"课程标准|课程教学大纲|教学大纲|课程大纲|人才培养方案|授课计划|教案|课件",
        "",
        name,
    )
    name = re.sub(r"[《》()（）【】\[\]_—\-]", " ", name)
    name = re.sub(r"\s+", " ", name).strip()

    return normalize_course_name(name)


def extract_section_by_keywords(text: str, keywords: List[str], max_chars: int) -> str:
    if not text:
        return ""

    lines = [line.strip() for line in text.splitlines() if line.strip()]
    stop_keywords = [
        "课程性质",
        "课程定位",
        "课程目标",
        "教学目标",
        "素质目标",
        "知识目标",
        "能力目标",
        "教学内容",
        "课程内容",
        "教学项目",
        "学习任务",
        "教学单元",
        "教学安排",
        "实施建议",
        "教学方法",
        "考核方式",
        "考核评价",
        "评价方式",
        "成绩评定",
        "考核办法",
        "参考资料",
    ]

    for index, line in enumerate(lines):
        if not any(keyword in line for keyword in keywords):
            continue

        section_lines = [line]
        for next_line in lines[index + 1 : index + 40]:
            is_next_section = any(keyword in next_line for keyword in stop_keywords)
            is_same_section = any(keyword in next_line for keyword in keywords)
            if section_lines and is_next_section and not is_same_section:
                break

            section_lines.append(next_line)
            if len("\n".join(section_lines)) >= max_chars:
                break

        return "\n".join(section_lines)

    return ""


def split_items(section_text: str) -> List[str]:
    if not section_text:
        return []

    invalid_titles = {
        "课程目标",
        "教学目标",
        "素质目标",
        "知识目标",
        "能力目标",
        "教学内容",
        "课程内容",
        "教学项目",
        "学习任务",
        "教学单元",
        "考核方式",
        "考核评价",
        "评价方式",
        "成绩评定",
        "考核办法",
    }

    items = []
    for line in section_text.splitlines():
        line = line.strip()
        if len(line) < 4:
            continue
        if line in invalid_titles:
            continue
        items.append(line)

    return items[:10]


def extract_course_standard_info(text: str) -> dict:
    if not text:
        return build_empty_course_standard_info()

    objective_section = extract_section_by_keywords(
        text,
        ["课程目标", "教学目标", "素质目标", "知识目标", "能力目标"],
        max_chars=1500,
    )
    content_section = extract_section_by_keywords(
        text,
        ["教学内容", "课程内容", "教学项目", "学习任务", "教学单元"],
        max_chars=2000,
    )
    assessment_section = extract_section_by_keywords(
        text,
        ["考核方式", "考核评价", "评价方式", "成绩评定", "考核办法"],
        max_chars=1500,
    )

    return {
        "course_objectives": split_items(objective_section),
        "teaching_content": split_items(content_section),
        "assessment_methods": split_items(assessment_section),
        "raw_sections": {
            "objective_section": objective_section,
            "content_section": content_section,
            "assessment_section": assessment_section,
        },
    }


def merge_course_standard_info(current_info: dict, new_info: dict) -> dict:
    merged_info = current_info or build_empty_course_standard_info()
    new_info = new_info or build_empty_course_standard_info()

    for key in ["course_objectives", "teaching_content", "assessment_methods"]:
        values = new_info.get(key) or []
        if values and not merged_info.get(key):
            merged_info[key] = values

    current_sections = merged_info.get("raw_sections") or {}
    new_sections = new_info.get("raw_sections") or {}
    for key in ["objective_section", "content_section", "assessment_section"]:
        if new_sections.get(key) and not current_sections.get(key):
            current_sections[key] = new_sections[key]
    merged_info["raw_sections"] = current_sections

    return merged_info


def generate_rule_based_audit(course_standard_info: dict) -> dict:
    if not isinstance(course_standard_info, dict):
        course_standard_info = build_empty_course_standard_info()

    course_objectives = course_standard_info.get("course_objectives") or []
    teaching_content = course_standard_info.get("teaching_content") or []
    assessment_methods = course_standard_info.get("assessment_methods") or []

    checks = [
        {
            "item": "课程目标",
            "status": "pass" if course_objectives else "warning",
            "message": (
                f"已识别到 {len(course_objectives)} 条课程目标。"
                if course_objectives
                else "未识别到明确的课程目标。"
            ),
        },
        {
            "item": "教学内容",
            "status": "pass" if teaching_content else "warning",
            "message": (
                f"已识别到 {len(teaching_content)} 条教学内容。"
                if teaching_content
                else "未识别到明确的教学内容。"
            ),
        },
        {
            "item": "考核方式",
            "status": "pass" if assessment_methods else "warning",
            "message": (
                f"已识别到 {len(assessment_methods)} 条考核方式。"
                if assessment_methods
                else "未识别到明确的考核方式。"
            ),
        },
    ]

    missing_items = [check["item"] for check in checks if check["status"] != "pass"]
    passed_count = len(checks) - len(missing_items)
    score = round(passed_count / len(checks) * 100)

    if missing_items:
        summary = "课程标准关键要素识别不完整：" + "、".join(missing_items)
    else:
        summary = "课程标准关键要素识别完整。"

    return {
        "audit_type": "rule_based",
        "score": score,
        "summary": summary,
        "checks": checks,
        "missing_items": missing_items,
    }


def call_deepseek_audit(course_name: str, audit_mode: str, output_type: str, course_standard_info: dict, syllabus_full_text: str = "") -> dict:
    """
    调用 DeepSeek 生成 AI 审核结果。
    如果调用失败，返回空字典，由后端继续使用规则审核结果兜底。
    """
    api_key = os.getenv("DEEPSEEK_API_KEY")
    print("开始调用 DeepSeek AI 审核")

    if not api_key:
        print("未设置 DEEPSEEK_API_KEY，跳过 AI 审核")
        return {}

    prompt = f"""
你是一名高职院校课程教学资料审核专家。请根据以下课程标准解析结果，进行教学资料一致性审核。

课程名称：{course_name}
审核模式：{audit_mode}
输出类型：{output_type}

课程目标：
{json.dumps(course_standard_info.get("course_objectives", []), ensure_ascii=False, indent=2)}

教学内容：
{json.dumps(course_standard_info.get("teaching_content", []), ensure_ascii=False, indent=2)}

考核方式：
{json.dumps(course_standard_info.get("assessment_methods", []), ensure_ascii=False, indent=2)}

课程标准原文片段：
{json.dumps(course_standard_info.get("raw_sections", {}), ensure_ascii=False, indent=2)}
课程标准全文节选：
{syllabus_full_text[:8000]}

请严格输出 JSON，不要输出任何解释性文字。
JSON 格式如下：
{{
  "score": 0-100的整数,
  "conclusion": "总体判断",
  "issues": [
    {{
      "level": "高/中/低",
      "type": "问题类型",
      "desc": "问题描述",
      "suggestion": "整改建议"
    }}
  ]
}}

要求：
0. 请以“课程标准全文节选”为主要依据，course_objectives、teaching_content、assessment_methods 仅作为辅助参考；
0. 不要评价“解析结果是否完整”，不要使用“解析结果中缺失”“系统未识别到”等表述；如果发现内容不足，请表述为“课程标准中未充分体现……”或“教学材料中未充分体现……”；
1. 不要使用与课程无关的表述；
2. 如果课程是艺术设计类，不要出现“译文质量”等翻译类表述；
3. 审核意见要具体、建设性，适合高职院校教学质量监控场景；
4. 如果没有明显问题，也要给出低风险复核建议。
"""

    try:
        response = requests.post(
            "https://api.deepseek.com/chat/completions",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            },
            json={
                "model": "deepseek-chat",
                "messages": [
                    {"role": "user", "content": prompt}
                ],
                "temperature": 0.2,
            },
            timeout=90,
        )

        if response.status_code != 200:
            print("DeepSeek 调用失败：", response.status_code, response.text)
            return {}

        data = response.json()
        content = data["choices"][0]["message"]["content"]

        # 防止模型返回 ```json 包裹
        content = content.strip()
        if content.startswith("```"):
            content = content.replace("```json", "").replace("```", "").strip()

        result = json.loads(content)

        if "score" not in result or "issues" not in result:
            print("DeepSeek 返回格式不完整：", result)
            return {}

        return result

    except Exception as e:
        print("DeepSeek 审核异常：", e)
        return {}
    
    
@app.get("/")
def read_root():
    return {"message": "AI Teaching Material Audit Backend is running."}


@app.post("/upload")
async def upload_files(
    program_files: List[UploadFile] = File(...),
    syllabus_files: List[UploadFile] = File(...),
    material_files: List[UploadFile] = File(...),
    audit_mode: str = Form("full"),
    output_type: str = Form("detailed"),
):
    saved_files = {
        "program_files": [],
        "syllabus_files": [],
        "material_files": [],
    }
    groups = {
        "program_files": program_files or [],
        "syllabus_files": syllabus_files or [],
        "material_files": material_files or [],
    }

    detected_course_name = ""
    syllabus_text_preview = ""
    syllabus_full_text = ""
    course_standard_info = build_empty_course_standard_info()

    for group_name, files in groups.items():
        group_dir = UPLOAD_DIR / group_name
        group_dir.mkdir(parents=True, exist_ok=True)

        for upload_file in files:
            filename = safe_filename(upload_file.filename)
            file_path = group_dir / filename
            file_record = {
                "filename": filename,
                "path": str(file_path),
                "content_type": upload_file.content_type,
            }

            try:
                with file_path.open("wb") as buffer:
                    shutil.copyfileobj(upload_file.file, buffer)
            except Exception as exc:
                file_record["error"] = f"save_failed: {exc}"
                saved_files[group_name].append(file_record)
                continue

            saved_files[group_name].append(file_record)

            if group_name != "syllabus_files":
                continue

            text = ""
            try:
                text = extract_text(file_path)
            except Exception as exc:
                file_record["error"] = f"parse_failed: {exc}"

            if text and not syllabus_text_preview:
                syllabus_text_preview = text[:500]
            if text and not syllabus_full_text:
                syllabus_full_text = text

            if text:
                extracted_info = extract_course_standard_info(text)
                course_standard_info = merge_course_standard_info(
                    course_standard_info,
                    extracted_info,
                )

            if not detected_course_name:
                detected_course_name = guess_course_name_from_text(text)

            if not detected_course_name:
                detected_course_name = guess_course_name_from_filename(filename)

    audit_result = generate_rule_based_audit(course_standard_info)
    ai_result = call_deepseek_audit(
        detected_course_name,
        audit_mode,
        output_type,
        course_standard_info,
        syllabus_full_text,
    )

    audit_source = "规则审核兜底"

    if ai_result:
        audit_result = ai_result
        audit_source = "DeepSeek AI审核"

    if audit_mode == "quick":
        audit_result["conclusion"] = "总体判断：本次为快速审核，重点检查课程标准与教学内容的基础匹配情况。"
        audit_result["issues"] = [
            issue
            for issue in audit_result.get("issues", [])
            if "教学内容" in issue.get("type", "") or "课程目标" in issue.get("type", "")
        ] or audit_result.get("issues", [])

    if audit_mode == "assessment":
        audit_result["conclusion"] = "总体判断：本次为考核一致性专项审核，重点检查考核方式是否支撑课程目标达成。"
        audit_result["issues"] = [
            issue
            for issue in audit_result.get("issues", [])
            if "考核" in issue.get("type", "") or "评价" in issue.get("type", "")
        ] or [
            {
                "level": "低",
                "type": "考核一致性需人工复核",
                "desc": "系统已识别课程标准中的考核相关信息，但仍建议进一步检查考核任务、评价标准与课程目标之间的对应关系。",
                "suggestion": "建议建立“课程目标—考核任务—评价标准”对应表，用于专项复核。",
            }
        ]

    return {
        "status": "success",
        "message": "文件上传成功",
        "detected_course_name": detected_course_name,
        "syllabus_text_preview": syllabus_text_preview,
        "course_standard_info": course_standard_info,
        "audit_result": audit_result,
        "audit_source": audit_source,
        "saved_files": saved_files,
    }
